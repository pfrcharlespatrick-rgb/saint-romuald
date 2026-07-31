# Nécessite l'installation préalable de : pip install python-docx

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


doc = docx.Document()

# Configuration des marges
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Titre principal
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = title.add_run(
    "Dictionnaire biographique & Synthèse historique\nde New Liverpool (Pages"
    " 15 à 60)"
)
run_title.font.name = "Georgia"
run_title.font.size = Pt(22)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Bleu marine royal

# Sous-titre
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = sub.add_run(
    "Dépouillement de l'ouvrage de Julie Doyon (Éditions Septentrion)"
)
run_sub.font.name = "Arial"
run_sub.font.size = Pt(12)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # Espacement

# Introduction
h1 = doc.add_heading(level=1)
r_h1 = h1.add_run(
    "1. Synthèse générale d'introduction : L'évolution du territoire"
)
r_h1.font.name = "Georgia"
r_h1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

p_intro = doc.add_paragraph()
p_intro.paragraph_format.line_spacing = 1.15
r_intro = p_intro.add_run(
    "L'étude des pages 15 à 60 retrace la transformation séculaire d'un espace"
    " littoral névralgique situé sur la rive sud du fleuve Saint-Laurent, à la"
    " jonction des rivières Chaudière et Etchemin. Autrefois halte saisonnière"
    " pour les Premières Nations, le secteur devient un foyer majeur du"
    " commerce du bois d'œuvre, de la construction navale et du transport"
    " fluvial au Bas-Canada sous l'impulsion de grandes dynasties marchandes."
)
r_intro.font.name = "Arial"
r_intro.font.size = Pt(10.5)

# Tableau récapitulatif
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
headers = ["Période / Thème", "Développements clés", "Acteurs principaux"]
for i, head in enumerate(headers):
    hdr_cells[i].text = head
    set_cell_background(hdr_cells[i], "1B365D")
    p = hdr_cells[i].paragraphs[0]
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Arial"

data = [
    (
        "Préhistoire & Régime français",
        (
            "Sites paléoindiens (Longwood), cabane d'Eustache Lambert (1651),"
            " moulin Strouds."
        ),
        "E. Lambert, F. Bissot, G. Couture, G.W. Strouds",
    ),
    (
        "Conquête & Régime anglais",
        (
            "Incendies de 1759-1760, vente de Lauzon à Murray (1765), essor sous"
            " H. Caldwell."
        ),
        "J. Wolfe, É. Charest III, J. Murray, H. Caldwell",
    ),
    (
        "Âge d'or du Bois (1806-1840)",
        (
            "Blocus napoléonien, chantiers navals Hamilton, moulins Ritchie,"
            " empire Price."
        ),
        "J. Caldwell, Frères Hamilton, Frères Ritchie, W. Price",
    ),
    (
        "Monopoles & Société (1840-1856)",
        (
            "Traversées à manège, régime des jetons, érection de la paroisse"
            " St-Romuald (1856)."
        ),
        "P. McLeod, P. Honorat, C. Poiré, Abbé P.-T. Sax",
    ),
]

for row_data in data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        p = row_cells[i].paragraphs[0]
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9.5)

doc.add_paragraph()  # Espacement

# Dictionnaire biographique
h2 = doc.add_heading(level=1)
r_h2 = h2.add_run("2. Dictionnaire biographique")
r_h2.font.name = "Georgia"
r_h2.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

h2_1 = doc.add_heading(level=2)
r_h2_1 = h2_1.add_run(
    "Magdeleine Delisle (1791-1873), pionnière de New Liverpool"
)
r_h2_1.font.name = "Georgia"
r_h2_1.font.bold = True
r_h2_1.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

p_bio_source = doc.add_paragraph()
r_bio_source = p_bio_source.add_run(
    "Source : Guy Saint-Hilaire, maître généalogiste agréé, « Magdeleine"
    " Delisle (1791-1873), pionnière de New Liverpool, prolifique souche des"
    " Hardy, Roberge, Gibson, Nolin, Pelletier et Forcade », L'Ancêtre,"
    " n° 292, vol. 37, automne 2010, p. 23-32."
)
r_bio_source.font.name = "Arial"
r_bio_source.font.size = Pt(9)
r_bio_source.font.italic = True
r_bio_source.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p_bio = doc.add_paragraph()
p_bio.paragraph_format.line_spacing = 1.15
r_bio = p_bio.add_run(
    "Née le 9 juillet 1791 à Saint-Jean-de-l'Île-d'Orléans, Magdeleine"
    " Delisle épouse Joseph Hardy en 1810 et s'établit avec lui à New"
    " Liverpool, où elle est abandonnée par son mari vers l'automne 1817"
    " après cinq enfants. Devenue propriétaire d'une résidence en 1823,"
    " elle a par la suite six autres enfants, de trois pères différents,"
    " dont un compagnonnage d'une quinzaine d'années avec le marchand"
    " Jean Pelletier. Des onze enfants qu'elle a eus, neuf se marient et"
    " laissent une abondante descendance sous les patronymes Hardy,"
    " Burroughs (devenu Roberge par alliance), Gibson, Nolin, Pelletier et"
    " Forcade. Magdeleine Delisle meurt le 14 février 1873 à Saint-Romuald,"
    " à l'âge de 81 ans, laissant au moins 86 petits-enfants identifiés."
    " Ses descendants Hardy, Roberge et Pelletier sont toujours"
    " représentés aujourd'hui à Saint-Romuald et dans le quartier de New"
    " Liverpool."
)
r_bio.font.name = "Arial"
r_bio.font.size = Pt(10.5)

# Tableau de la descendance
table2 = doc.add_table(rows=1, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ["Enfant", "Conjoint", "Mariage", "Postérité"]
hdr_cells2 = table2.rows[0].cells
for i, head in enumerate(headers2):
    hdr_cells2[i].text = head
    set_cell_background(hdr_cells2[i], "1B365D")
    p = hdr_cells2[i].paragraphs[0]
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Arial"

data2 = [
    ("Joseph Hardy (fils)", "Aveline Fortin", "14 fév. 1848, Plessisville", "8 enfants"),
    ("David Hardy", "Rose-de-Lima Roberge", "5 oct. 1841, St-Jean-Chrysostome", "12 enfants — souche des Hardy de Saint-Romuald"),
    ("Henriette Hardy", "Télesphore Marchand", "1er fév. 1848, Plessisville", "6 enfants"),
    ("Élisabeth (Isabelle) Burroughs", "Hubert Roberge", "21 janv. 1840, St-Jean-Chrysostome", "14 enfants"),
    ("James Gibson", "Marie-Louise McReady", "22 nov. 1842, St-Jean-Chrysostome", "11 enfants"),
    ("Éléonore Pelletier", "Étienne Nolin, puis Victor Girouard", "13 sept. 1842 / 22 fév. 1881", "13 enfants"),
    ("Édouard Pelletier", "Céleste Roberge", "18 sept. 1848, St-Jean-Chrysostome", "11 enfants"),
    ("Caroline Pelletier", "Joseph Forcade, puis Norbert Gosselin (fils)", "18 sept. 1848 / 9 janv. 1872", "6 enfants"),
    ("Eusèbe Pelletier", "Marie-Émilie McReady, puis Léda Couillard, puis Eusébie Houle", "20 fév. 1860, Saint-Romuald", "5 enfants"),
]

for row_data in data2:
    row_cells = table2.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        p = row_cells[i].paragraphs[0]
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9)

# Enregistrement du fichier
doc.save("Synthese_Histoire_New_Liverpool.docx")
print("Document Word généré avec succès : Synthese_Histoire_New_Liverpool.docx")
